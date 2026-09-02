#!/usr/bin/env python3
"""
Extract and analyze Knowledge Requirements from DOCX file
9/1/2026 - rt
"""

import sys
from pathlib import Path

# Try to use python-docx library if available, otherwise try alternative
try:
    from docx import Document
    
    docx_path = Path(r"C:\Users\Ron Treleaven\Drone_SOP\Advanced_RPAS_Study\docs\Knowledge_Requirements_Advanced.docx")
    
    doc = Document(docx_path)
    
    print("=" * 100)
    print("KNOWLEDGE REQUIREMENTS AND LINKS ANALYSIS")
    print("=" * 100)
    print()
    
    # Extract all text from document
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    
    print("\n" + "=" * 100)
    print("TABLES IN DOCUMENT:")
    print("=" * 100)
    
    # Extract tables
    for i, table in enumerate(doc.tables):
        print(f"\n[TABLE {i+1}]")
        for j, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            print(" | ".join(cells))
    
except ImportError:
    print("python-docx not available. Attempting alternative extraction...")
    
    # Alternative: Try to extract XML from DOCX (it's a ZIP file)
    import zipfile
    import xml.etree.ElementTree as ET
    
    docx_path = Path(r"C:\Users\Ron Treleaven\Drone_SOP\Advanced_RPAS_Study\docs\Knowledge_Requirements_Advanced.docx")
    
    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            xml_content = zip_ref.read('word/document.xml')
            
            # Parse XML
            root = ET.fromstring(xml_content)
            
            # Namespace for Word documents
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
            
            # Extract paragraphs
            print("=" * 100)
            print("KNOWLEDGE REQUIREMENTS AND LINKS ANALYSIS")
            print("=" * 100)
            print()
            
            for para in root.findall('.//w:p', namespaces):
                text_elements = para.findall('.//w:t', namespaces)
                para_text = ''.join([t.text for t in text_elements if t.text])
                if para_text.strip():
                    print(para_text)
                    
    except Exception as e:
        print(f"Error reading document: {e}")
        print("Please ensure the file path is correct and the document is readable.")
